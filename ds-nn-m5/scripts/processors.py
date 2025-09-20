# processor.py
from docx import Document
import fitz  # PyMuPDF
from typing import List, Dict
from scripts.models import TextElement, PageContent, FormatIssue

class DocumentProcessor:
    def __init__(self):
        self.font_sizes = {
            'Heading 1': 13,
            'Heading 2': 12,
            'Heading 3': 11,
            'Normal': 10
        }
        self.pdf_tolerance = 0.2

    def _get_heading_level(self, style_name: str) -> str:
        """Determine heading level from style name"""
        style_lower = style_name.lower()
        if 'heading 1' in style_lower or 'h1' in style_lower:
            return 'Heading 1'
        elif 'heading 2' in style_lower or 'h2' in style_lower:
            return 'Heading 2'
        elif 'heading 3' in style_lower or 'h3' in style_lower:
            return 'Heading 3'
        return 'Normal'

    def process_docx(self, file_path: str) -> List[FormatIssue]:
        doc = Document(file_path)
        issues = []
        
        # Calculate pages based on sections
        total_pages = self._calculate_pages(doc)
        current_page = 1
        chars_per_page = 3000
        char_count = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                char_count += len(para.text)
                current_page = min(total_pages, (char_count // chars_per_page) + 1)
                
                if para.runs:
                    font_size = None
                    for run in para.runs:
                        if run.font.size:
                            font_size = run.font.size.pt
                            break
                    
                    # Get the normalized style name
                    style = self._get_heading_level(para.style.name)
                    
                    if not font_size:
                        font_size = self.font_sizes.get(style, self.font_sizes['Normal'])
                    
                    # Check for format issues
                    expected_size = self.font_sizes.get(style)
                    if expected_size and font_size != expected_size:
                        issues.append(FormatIssue(
                            page=current_page,
                            text=para.text[:100],
                            current_size=font_size,
                            expected_size=expected_size,
                            style=style
                        ))
        
        return issues

    def _calculate_pages(self, doc) -> int:
        """Estimate total pages in the document"""
        total_chars = sum(len(paragraph.text) for paragraph in doc.paragraphs)
        chars_per_page = 3000  # Approximate characters per page
        return max(1, (total_chars // chars_per_page) + 1)

    def process_pdf(self, file_path: str) -> List[FormatIssue]:
        pdf = fitz.open(file_path)
        issues = []
        
        try:
            # Process each page
            for page_num in range(pdf.page_count):
                page = pdf[page_num]
                blocks = page.get_text("dict")["blocks"]
                
                # Process text blocks on page
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if text:
                                    font_size = round(span["size"], 1)
                                    style = self._determine_style(span["font"], font_size)
                                    
                                    # Check format with tolerance
                                    expected_size = self.font_sizes.get(style)
                                    if expected_size and abs(font_size - expected_size) > self.pdf_tolerance:
                                        issues.append(FormatIssue(
                                            page=page_num + 1,
                                            text=text[:100],
                                            current_size=font_size,
                                            expected_size=expected_size,
                                            style=style
                                        ))
            
            pdf.close()
            return issues
            
        except Exception as e:
            if pdf:
                pdf.close()
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def _determine_style(self, font_name: str, font_size: float) -> str:
        """Determine text style based on font properties"""
        is_bold = "bold" in font_name.lower() or "heavy" in font_name.lower()
        
        if is_bold:
            if font_size >= 13 - self.pdf_tolerance:
                return 'Heading 1'
            elif font_size >= 12 - self.pdf_tolerance:
                return 'Heading 2'
            elif font_size >= 11 - self.pdf_tolerance:
                return 'Heading 3'
        return 'Normal'