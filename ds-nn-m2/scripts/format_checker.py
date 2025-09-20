from docx import Document
from docx.shared import Inches, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt
from typing import List, Dict, Tuple, Optional
from common.logs import logger
from common.s3_operations import S3Helper
import re, os
from PyPDF2 import PdfReader
from typing import List, Dict, Optional, Any
import re
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar, LTAnno, LTTextBox
import fitz  # PyMuPDF
import subprocess
from lxml import etree

TMP_DIR = '/tmp'
if not os.path.exists(TMP_DIR):
    os.makedirs(TMP_DIR)

class DocumentFormatReviewer:
    def __init__(self, doc_path: str):
        """Initialize with path to Word document and logger"""
        try:
            logger.info(f"Opening document: {doc_path}")
            if doc_path.startswith('s3://'):
                doc_path = self.download_s3_file(doc_path)
            self.document = Document(doc_path)
            self.file_path = doc_path
            self.heading_errors = []
            self.margin_errors = []
            self.bullet_errors = []
            logger.info("Document loaded successfully")
        except Exception as e:
            logger.error(f"Failed to open document: {str(e)}")
            raise Exception(f"Failed to open document: {str(e)}")

    def download_s3_file(self, word_path):
        s3_bucket = word_path.split('/')[2]
        s3_helper = S3Helper(s3_bucket)
        s3_key = '/'.join(word_path.split('/')[3:])
        local_file_path = os.path.join(TMP_DIR, os.path.basename(s3_key))
                # Download the file from S3
        s3_helper.download_file_from_s3(s3_key, local_file_path)
        return local_file_path

    # def _get_file_type(self) -> str:
    #     """Determine file type from extension"""
    #     ext = os.path.splitext(self.file_path)[1].lower()
    #     if ext in ['.doc', '.docx']:
    #         return 'word'
    #     elif ext == '.pdf':
    #         return 'pdf'
    #     else:
    #         raise ValueError(f"Unsupported file type: {ext}")

    # def _init_document(self):
    #     """Initialize document based on file type"""
    #     try:
    #         logger.info(f"Opening {self.file_type.upper()} document: {self.file_path}")
    #         if self.file_type == 'word':
    #             self.doc = Document(self.file_path)
    #         else:
    #             self.doc = fitz.open(self.file_path)
    #     except Exception as e:
    #         logger.error(f"Failed to open document: {str(e)}")
    #         raise

    def convert_to_pdf(self) -> str:
        """Convert document to PDF using LibreOffice"""
        output_dir = os.path.dirname(self.file_path)
        pdf_file = os.path.join(output_dir, os.path.splitext(os.path.basename(self.file_path))[0] + '.pdf')
        
        try:
            command = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                self.file_path
            ]
            subprocess.run(command, check=True)
            return pdf_file
        except subprocess.CalledProcessError as e:
            logger.error(f"PDF conversion failed: {str(e)}")
            raise

    def extract_word_headings(self) -> List:
        """Extract headings from Word document"""
        headings = []
        for para in self.document.paragraphs:
            if 'Heading' in para.style.name or 'Header' in para.style.name:
                if para.text.strip():
                    heading = (
                        para.text.strip(),
                        para.style.name
                    )
                    headings.append(heading)
        return headings

    def preprocess_pdf_text(self,text):
        """
        Preprocesses PDF text to merge logical lines and handle formatting quirks.
        """
        lines = text.splitlines()
        processed_lines = []
        buffer_line = ""

        for line in lines:
            stripped_line = line.strip()

            if not stripped_line:  # Empty line, treat as a paragraph break
                if buffer_line:
                    processed_lines.append(buffer_line)
                    buffer_line = ""
                continue

            # Heuristic: Check if the line is a continuation of the previous one
            if buffer_line and not stripped_line[0].isupper() and not re.match(r'^\d+(\.\d+)*', stripped_line):
                buffer_line += f" {stripped_line}"
            else:
                if buffer_line:  # Add previous buffer to processed lines
                    processed_lines.append(buffer_line)
                buffer_line = stripped_line  # Start a new buffer

        if buffer_line:  # Add the last buffer line
            processed_lines.append(buffer_line)

        return "\n".join(processed_lines)

    def extract_heading_numbers_from_pdf(self, pdf_file: str, docx_headings: List) -> Dict[str, str]:
        """Extract heading numbers from PDF"""
        doc = None
        try:
            doc = fitz.open(pdf_file)
            full_text = ""
            for page in doc:
                full_text += page.get_text("text") + "\n"
            
            full_text = self.preprocess_pdf_text(full_text)
            # full_text = re.sub(r'\n(?!\n)', ' ', full_text)
            print(full_text)
            matched_headings = {}
            for heading in docx_headings:
                if heading[0]:
                    # heading_parts = [re.escape(part.strip()) for part in heading[0].split('\n')]
# Join the parts with a flexible whitespace/newline matcher
                    # heading_regex = r'\s*(?:\r?\n|\s)*'.join(heading_parts)
                    heading_regex = re.escape(heading[0]).replace(r"\ ", r"\s*(?:\r?\n|\s)*")
                    # Create the combined pattern that matches the number and multi-line heading
                    combined_regex = rf"(?<!\S)\s*(\d+(?:\.\d+)*\.?)\s*(?:\r?\n|\s)*({heading_regex})"
                    print(combined_regex)
                    if re.search(combined_regex, full_text, re.DOTALL) is not None:
                        matches = re.findall(combined_regex, full_text,re.DOTALL)
                        # print(matchess)
                        # for match in matches:
                        #     print(match)
                        #     print(match.group(1))
                        # print(heading[0])
                        if matches:
                            # print(1)
                            # print(len(matches))
                            for match in matches:
                                # print('NUMBER')
                                print(match)
                                if match[0]:
                                    matched_headings[heading[0]] = match[0]
                                else:
                                    if not matched_headings.get(heading[0], None):
                                        matched_headings[heading[0]] = None
                                # print('HEADING')
                                # print(match.group(2)[0:5]
                    else:
                        # print(2)
                        matched_headings[heading[0]] = None
                    # print(matched_headings[heading[0]])
                    print('---------------------------------------------')
                    # if matchess:
                    #     for matches in matchess:
                    #         if matches[-1][1] == '':
                    #             matched_headings[heading[0]] = matches[-1][0]
                    #         else:
                    #             matched_headings[heading[0]] = matches[-1][1].strip()
                    # else:
                    #     matched_headings[heading[0]] = None
            print(matched_headings)
            return matched_headings
        except Exception as e:
            logger.error(str(f'{e} ')+'[extract_heading_numbers_from_pdf] [scripts/format_checker.py:127]')
            raise Exception(f'{e}')

    def normalize_number(self, number: Optional[str]) -> Optional[Tuple[int, ...]]:
        """Normalize heading numbers for comparison"""
        if not number or number.strip() == '':
            return None
        number = number.strip().rstrip('.')
        try:
            parts = [part for part in number.split('.') if part]
            return tuple(map(int, parts))
        except (ValueError, AttributeError):
            return None

    def check_heading_hierarchy(self, heading_numbers: Dict[str, str]) -> List[Dict[str, Any]]:
        """Validate heading hierarchy and numbering"""
        errors = []
        used_numbers = set()
        seen_prefixes = set()
        last_number_at_level = {}  # Track last number seen at each level with prefix
        max_section_seen = 0  # Track highest top-level section number
        sequence_broken = False  # Flag to track if a higher section number was seen
        
        # First pass: Check for missing and invalid numbers
        numbered_headings = []
        for heading_text, number in heading_numbers.items():
            if not number:
                errors.append({
                    "error_type": "Missing number",
                    "heading": heading_text,
                    "number": None,
                    "details": None
                })
                continue
                
            norm_number = self.normalize_number(number)
            if not norm_number:
                errors.append({
                    "error_type": "Invalid format",
                    "heading": heading_text,
                    "number": number,
                    "details": None
                })
            else:
                numbered_headings.append((heading_text, number, norm_number))
        
        # Check each heading
        for heading_text, number, norm_number in numbered_headings:
            # Check for reused numbers
            if norm_number in used_numbers:
                errors.append({
                    "error_type": "Number reuse",
                    "heading": heading_text,
                    "number": number,
                    "details": "Number already used"
                })
                continue
            used_numbers.add(norm_number)
            
            # Handle top-level sections
            if len(norm_number) == 1:
                if norm_number[0] > max_section_seen:
                    max_section_seen = norm_number[0]
                    sequence_broken = True
                elif norm_number[0] == max_section_seen:
                    errors.append({
                        "error_type": "Number reuse",
                        "heading": heading_text,
                        "number": number,
                        "details": "Section number already used"
                    })
            
            # Handle subsections
            else:
                current_level = len(norm_number)
                prefix = norm_number[:-1]
                current_num = norm_number[-1]
                
                # Check for incorrect sequences after higher section number
                if sequence_broken and norm_number[0] < max_section_seen:
                    errors.append({
                        "error_type": "Incorrect sequence",
                        "heading": heading_text,
                        "number": number,
                        "details": f"Cannot use section {norm_number[0]} after section {max_section_seen}"
                    })
                    continue
                
                # Check for missing parent sections
                for i in range(1, current_level):
                    intermediate_prefix = norm_number[:i]
                    if intermediate_prefix not in seen_prefixes:
                        errors.append({
                            "error_type": "Incorrect sequence",
                            "heading": heading_text,
                            "number": number,
                            "details": f"Missing parent section {'.'.join(map(str, intermediate_prefix))}"
                        })
                
                # Check sequence within level
                if prefix not in last_number_at_level:
                    # First number at this level should be 1
                    if current_num != 1:
                        errors.append({
                            "error_type": "Incorrect sequence",
                            "heading": heading_text,
                            "number": number,
                            "details": f"First number at this level should be 1"
                        })
                else:
                    # Check if number follows sequence
                    last_num = last_number_at_level[prefix]
                    if current_num != last_num + 1:
                        errors.append({
                            "error_type": "Incorrect sequence",
                            "heading": heading_text,
                            "number": number,
                            "details": f"Incorrect number sequence after {'.'.join(map(str, prefix))}.{last_num}"
                        })
                
                last_number_at_level[prefix] = current_num
            
            # Add all prefixes of current number to seen prefixes
            for i in range(1, len(norm_number) + 1):
                seen_prefixes.add(norm_number[:i])
        
        return errors

    def review_document_headings(self) -> List[str]:
        """Perform complete document review"""
        try:
            # if self.file_type == 'word':
                # Extract headings from Word
            docx_headings = self.extract_word_headings()
                
                # Convert to PDF and extract numbers
            pdf_file = self.convert_to_pdf()
            heading_numbers = self.extract_heading_numbers_from_pdf(pdf_file, docx_headings)
                
                # Check hierarchy
            errors = self.check_heading_hierarchy(heading_numbers)
            return errors
            # else:
            #     logger.error("PDF review not implemented")
            #     return ["PDF review not implemented"]
                
        except Exception as e:
            logger.error(f"Document review failed: {str(e)}")
            raise Exception(f'{e}')
    # def get_paragraph_level(self, paragraph) -> Optional[int]:
    #     try:
    #         if paragraph.style and paragraph.style.name.startswith('Heading'):
    #             return int(paragraph.style.name[-1])
    #         return None
    #     except (AttributeError, ValueError) as e:
    #         logger.error(f"Error getting paragraph level: {str(e)}")
    #         return None

    # def check_heading_numbering(self) -> List[str]:
    #     logger.info("Starting heading numbering check")
    #     errors = []
    #     current_levels = [0] * 9
        
    #     try:
    #         for idx, paragraph in enumerate(self.document.paragraphs, 1):
    #             level = self.get_paragraph_level(paragraph)
    #             if level is not None:
    #                 logger.info(f"Checking heading at paragraph {idx}, level {level}")
    #                 level -= 1
                    
    #                 for i in range(level):
    #                     if current_levels[i] == 0:
    #                         error_msg = f"Heading level {level + 1} found before level {i + 1}"
    #                         logger.error(error_msg)
    #                         errors.append(error_msg)
                    
    #                 current_levels[level] += 1
    #                 for i in range(level + 1, 9):
    #                     current_levels[i] = 0
                    
    #                 expected_number = '.'.join(str(n) for n in current_levels[:level + 1] if n > 0)
    #                 if not paragraph.text.startswith(expected_number):
    #                     error_msg = f"Incorrect heading numbering: '{paragraph.text}' should start with '{expected_number}'"
    #                     logger.error(error_msg)
    #                     errors.append(error_msg)
    #                 else:
    #                     logger.info(f"Heading numbering correct: {paragraph.text}")
                        
    #     except Exception as e:
    #         error_msg = f"Error in heading numbering check: {str(e)}"
    #         logger.error(error_msg)
    #         raise
            
    #     logger.info(f"Completed heading numbering check. Found {len(errors)} errors")
    #     return errors

    # def check_heading_levels(self) -> List[str]:
    #     logger.info("Starting heading levels check")
    #     errors = []
    #     prev_level = 0
        
    #     try:
    #         for idx, paragraph in enumerate(self.document.paragraphs, 1):
    #             level = self.get_paragraph_level(paragraph)
    #             if level is not None:
    #                 logger.info(f"Checking heading level at paragraph {idx}: Level {level}")
    #                 if level > prev_level + 1:
    #                     error_msg = f"Invalid heading level jump: {prev_level} to {level} at '{paragraph.text}'"
    #                     logger.error(error_msg)
    #                     errors.append(error_msg)
    #                 prev_level = level
                    
    #     except Exception as e:
    #         error_msg = f"Error in heading levels check: {str(e)}"
    #         logger.error(error_msg)
    #         raise
            
    #     logger.info(f"Completed heading levels check. Found {len(errors)} errors")
    #     return errors

    # def get_margin_value(self, section_margin) -> Optional[float]:
    #     """Safely extract margin value in inches"""
    #     try:
    #         # Get the underlying twips value
    #         margin_twips = section_margin._element.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
    #         if margin_twips is None:
    #             logger.error("Could not find margin value in document")
    #             return None
                
    #         # Convert twips to inches (1 inch = 1440 twips)
    #         margin_inches = float(margin_twips) / 1440
    #         logger.info(f"Converted margin from {margin_twips} twips to {margin_inches:.2f} inches")
    #         return margin_inches
            
    #     except (AttributeError, KeyError, ValueError, TypeError) as e:
    #         logger.error(f"Error extracting margin value: {str(e)}")
    #         return None

    def check_page_margins(self) -> List[str]:
        logger.info("Starting page margins check")
        
        try:
            # distinct_margins = {}  # Dictionary to track distinct margins and the corresponding sections
            document_path = self.file_path.replace('.docx', '.pdf')
            pdf_reviewer = PDFFormatReviewer(document_path)
            result = pdf_reviewer.check_page_margins()
            # document = self.document
            # sections = document.sections

            # for i, section in enumerate(sections, 1):
            #     margins = {
            #         'top': section.top_margin.inches,
            #         'bottom': section.bottom_margin.inches,
            #         'left': section.left_margin.inches,
            #         'right': section.right_margin.inches
            #     }
            #     print(section)
            #     # Track distinct margins
            #     margin_tuple = tuple(sorted(margins.values()))  # Sort margins to avoid permutations being treated as distinct
            #     if margin_tuple not in distinct_margins:
            #         distinct_margins[margin_tuple] = []
            #     distinct_margins[margin_tuple].append(i)
            
            # # Display all distinct margins and the sections they belong to
            # result = []
            # for margin, pages in distinct_margins.items():
            #     result.append(f"Distinct margin: {margin} on pages {', '.join(map(str, pages))}")

            # # print(result)
                    
        except Exception as e:
            error_msg = f"Error in page margins check: {str(e)}"
            logger.error(error_msg)
            raise
            
        return result

    def twips_to_inches(self, twips_obj) -> Optional[float]:
        """Convert a twips measurement to inches"""
        try:
            # Try to get the raw twips value
            if hasattr(twips_obj, '_element'):
                twips = int(twips_obj._element.attrib.get('w:val', 0))
            else:
                # If no _element attribute, try direct conversion
                twips = int(str(twips_obj))
                
            # Convert to inches (1 inch = 1440 twips)
            inches = twips / 1440.0
            return round(inches, 2)
            
        except (AttributeError, ValueError, TypeError) as e:
            logger.error(f"Error converting twips to inches: {str(e)}")
            return None

    # def check_bullet_points(self) -> List[Dict[str, Any]]:
    #     """
    #     Comprehensive bullet point detection across all levels.
        
    #     Detects bullet points using a unified set of symbols.
    #     """
    #     logger.info("Starting comprehensive bullet point check")
    #     errors = []

    #     # Predefined acceptable bullet patterns
    #     acceptable_bullet_patterns = {
    #         1: '•',    # Solid dot
    #         2: '○',    # Hollow circle
    #         3: '■'     # Solid square
    #     }

    #     # Compiled list of all possible bullet point symbols
    #     ALL_BULLET_SYMBOLS = [
    #         # Level 1 symbols
    #         '•', '-', '*', '●', '►', '>', '➤', '→', '⁃', 
    #         # Level 2 symbols
    #         '○', '◯', '◦', '▫️', '◇', '▷', '»', '⦁', 
    #         # Level 3 symbols
    #         '■', '□', '▪️', '▫️', '▣', '▤', '➔', '⬧'
    #     ]

    #     try:
    #         for idx, paragraph in enumerate(self.document.paragraphs, 1):
    #             # Skip paragraphs that are not list paragraphs
    #             # if paragraph.style.name != 'List Paragraph':
    #             #     continue

    #             # Clean and prepare text
    #             cleaned_text = paragraph.text.strip()
    #             if not cleaned_text:
    #                 continue

    #             # First, detect if the paragraph starts with any bullet symbol
    #             first_char = cleaned_text[0]
    #             print(cleaned_text)
                
    #             # Check if the first character is a bullet symbol
    #             if first_char in ALL_BULLET_SYMBOLS:
    #                 # Calculate indent-based level
    #                 try:
    #                     indent = paragraph.paragraph_format.left_indent
    #                     indent_pt = float(indent.pt) if hasattr(indent, 'pt') else 0
    #                     level = min(3, max(1, int(indent_pt / 36)))
    #                 except Exception:
    #                     level = 1  # Default to level 1 if indent calculation fails

    #                 # Check if the bullet matches the expected symbol for its level
    #                 expected_bullet = acceptable_bullet_patterns.get(level)
                    
    #                 # If the detected symbol doesn't match the expected bullet for its level
    #                 if first_char != expected_bullet:
    #                     errors.append({
    #                         'paragraph_index': idx,
    #                         'detected_level': level,
    #                         'actual_bullet': first_char,
    #                         'expected_bullet': expected_bullet,
    #                         'full_text': cleaned_text,
    #                         'error_type': 'Incorrect Bullet Symbol'
    #                     })

    #     except Exception as global_error:
    #         logger.error(f"Global error in bullet point detection: {global_error}")
    #         errors.append({
    #             'error_type': 'Global Detection Failure',
    #             'error_message': str(global_error)
    #         })

    #     logger.info(f"Completed comprehensive bullet points check. Found {len(errors)} potential issues")
    #     return errors

    def check_bullet_points(self):
        # Open the .docx file
        doc = self.document
        rev = PDFFormatReviewer(self.file_path.replace('.docx', '.pdf'))
        return rev.check_bullet_points()
        
        # Initialize variables
        bullet_points = []
        indent_levels = []  # To track unique indentation levels for non-numbered bullets

        for paragraph in doc.paragraphs:
            # Access the raw XML of the paragraph
            xml = paragraph._p.xml
            tree = etree.fromstring(xml)
            
            # Check if the paragraph belongs to a list (bullet or numbered)
            numPr = tree.find('.//w:numPr', namespaces=paragraph._element.nsmap)
            
            if numPr is not None:
                # Get the list level
                ilvl = numPr.find('.//w:ilvl', namespaces=paragraph._element.nsmap)
                ilvl_value = int(ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")) if ilvl is not None else 0

                # Check for bullet symbol
                bullet_symbol = None
                bullet = tree.find('.//w:buChar', namespaces=paragraph._element.nsmap)
                if bullet is not None:
                    bullet_symbol = bullet.text.strip()
                
                # Add the bullet point text, level, and symbol
                bullet_points.append({
                    "text": paragraph.text.strip(),
                    "level": ilvl_value + 1,  # Convert to 1-based level
                    "symbol": bullet_symbol if bullet_symbol else "Unknown"
                })

        # Group by levels
        grouped_points = {1: [], 2: [], 3: []}
        for point in bullet_points:
            if point["level"] in grouped_points:
                grouped_points[point["level"]].append(point)
            else:
                grouped_points[1].append(point)  # Default to Level 1 if unexpected
        
        return grouped_points

    def review_document(self) -> Dict[str, List[str]]:
        logger.info("Starting complete document review")
        try:
            results = {
                'heading_numbering': self.review_document_headings(),
                'page_margins': self.check_page_margins(),
                'bullet_points': self.check_bullet_points()
            }
            
            total_errors = sum(len(issues) for issues in results.values())
            logger.info(f"Document review completed. Total errors found: {total_errors}")
            return results
            
        except Exception as e:
            error_msg = f"Failed to complete document review: {str(e)}"
            logger.error(error_msg)
            raise

# Example usage
# if __name__ == "__main__":
#     # Initialize with document path
#     reviewer = DocumentFormatReviewer("/home/yash-stride/Downloads/Sample file1_eQUAL.docx")
    
#     # Define expected margins (in inches)
#     margins = {
#         'top': 1.0,
#         'bottom': 1.0,
#         'left': 1.25,
#         'right': 1.25
#     }
    
#     # Perform review
#     results = reviewer.review_document(margins)
    
#     # Print results
#     for category, issues in results.items():
#         print(f"\n{category.replace('_', ' ').title()}:")
#         if issues:
#             for issue in issues:
#                 print(f"- {issue}")
#         else:
#             print("No issues found")


class PDFFormatReviewer:
    def __init__(self, pdf_path: str):
        """Initialize with path to PDF document and logger"""
        try:
            logger.info(f"Opening PDF: {pdf_path}")
            if pdf_path.startswith('s3://'):
                pdf_path = self.download_s3_file(pdf_path)
            self.pdf_reader = PdfReader(pdf_path)
            self.doc = fitz.open(pdf_path)  # PyMuPDF document
            logger.info(f"PDF loaded successfully. Total pages: {len(self.pdf_reader.pages)}")
        except Exception as e:
            logger.error(f"Failed to open PDF: {str(e)}")
            raise
    
    def download_s3_file(self, word_path):
        s3_bucket = word_path.split('/')[2]
        s3_helper = S3Helper(s3_bucket)
        s3_key = '/'.join(word_path.split('/')[3:])
        local_file_path = os.path.join(TMP_DIR, os.path.basename(s3_key))
                # Download the file from S3
        s3_helper.download_file_from_s3(s3_key, local_file_path)
        return local_file_path
    
    def get_heading_level(self, text: str, font_size: float) -> Optional[int]:
        """Determine heading level based on font size and formatting"""
        try:
            # Common heading patterns (e.g., "1.", "1.1", "1.1.1")
            heading_pattern = r'^\d+(\.\d+)*\s+'
            
            if re.match(heading_pattern, text):
                # Determine level based on number of dots + 1
                dots = text.split()[0].count('.')
                level = dots + 1
                logger.info(f"Detected heading level {level} for text: {text}")
                return level
            return None
        except Exception as e:
            logger.error(f"Error determining heading level: {str(e)}")
            return None

    def check_heading_numbering(self) -> List[str]:
        """Check if headings are numbered correctly and sequentially"""
        logger.info("Starting heading numbering check")
        errors = []
        current_levels = [0] * 9  # Track numbering for up to 9 levels
        
        try:
            for page_num, page in enumerate(self.doc, 1):
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                font_size = span["size"]
                                
                                level = self.get_heading_level(text, font_size)
                                if level is not None:
                                    logger.info(f"Checking heading on page {page_num}: {text}")
                                    level -= 1  # Convert to 0-based index
                                    
                                    # Check if higher level headings exist
                                    for i in range(level):
                                        if current_levels[i] == 0:
                                            error_msg = f"Page {page_num}: Heading level {level + 1} found before level {i + 1}"
                                            logger.error(error_msg)
                                            errors.append(error_msg)
                                    
                                    # Update numbering
                                    current_levels[level] += 1
                                    for i in range(level + 1, 9):
                                        current_levels[i] = 0
                                    
                                    # Check numbering format
                                    expected_number = '.'.join(str(n) for n in current_levels[:level + 1] if n > 0)
                                    if not text.startswith(expected_number):
                                        error_msg = f"Page {page_num}: Incorrect heading numbering: '{text}' should start with '{expected_number}'"
                                        logger.error(error_msg)
                                        errors.append(error_msg)
                                    
        except Exception as e:
            error_msg = f"Error in heading numbering check: {str(e)}"
            logger.error(error_msg)
            raise
            
        logger.info(f"Completed heading numbering check. Found {len(errors)} errors")
        return errors

    def check_page_margins(self) -> List[str]:
        """Check if page margins meet the requirements"""
        logger.info("Starting page margins check")
        distinct_margins = {}
        
        try:
            for i in range(self.doc.page_count):
                page = self.doc.load_page(i)
                
                # Get page dimensions (in points, 1 inch = 72 points)
                page_rect = page.rect
                page_width = page_rect.width
                page_height = page_rect.height
                
                # Define the margins based on page content (or assume default margins)
                # Assuming the page content starts at a fixed position and doesn't have a header/footer
                blocks = page.get_text("dict")["blocks"]
                        
                        # Track content boundaries
                left_margin = page_width
                right_margin = 0
                top_margin = page_height
                bottom_margin = 0
                        
                        # Check each block's position
                for block in blocks:
                    if "bbox" in block:
                        x0, y0, x1, y1 = block["bbox"]
                                
                                # Update content boundaries
                        left_margin = min(left_margin, x0)
                        right_margin = max(right_margin, x1)
                        top_margin = min(top_margin, y0)
                        bottom_margin = max(bottom_margin, y1)
                
                
                # Convert to inches (1 point = 1/72 inch)

                # Convert to inches (1 point = 1/72 inch)
                margins = {
                    'top': round(top_margin / 72, 2),  # Convert points to inches
                    'bottom': round(bottom_margin / 72, 2),
                    'left': round(left_margin / 72, 2),
                    'right': round(right_margin / 72, 2)
                }

                # Track distinct margins
                margin_tuple = tuple(margins.values())  # Sort margins to avoid permutations being treated as distinct
                if margin_tuple not in distinct_margins:
                    distinct_margins[margin_tuple] = []
                distinct_margins[margin_tuple].append(i + 1)  # Page numbers are 1-based


            distinct_margins = dict(sorted(distinct_margins.items(), key=lambda item: len(item[1])))  
            print(distinct_margins)

            # Display all distinct margins and the pages they belong to
            result = []
            for margin, pages in distinct_margins.items():
                result.append({"Margin": margin, "Pages": pages})

            # print(result)
                        
        except Exception as e:
            error_msg = f"Error in page margins check: {str(e)}"
            logger.error(error_msg)
            raise
            
        return result

    def check_bullet_points(self) -> List[dict]:
        """
        Check bullet point formatting and levels based on parent-child relationships.
        Bullet point hierarchy:
        - Level 1: Solid circle (•)
        - Level 2: Hollow circle (○)
        - Level 3: Solid square (■)
        """
        logger.info("Starting bullet points check")
        errors = []

        # Define bullet patterns for levels
        bullet_patterns = {
            1: ['•', '●'],  # Level 1: Solid circle
            2: ['○', '◦', 'o'],  # Level 2: Hollow circle
            3: ['■', '▪️']   # Level 3: Solid square
        }

        # Treat these as recognized bullet points
        recognized_bullet_symbols = [
            '•', '-', '*', '●', '►', '>', '➤', '→', '⁃',  # Common Level 1 bullets
            '○', '◯', '◦', '▫️', '◇', '▷', '»', '⦁',      # Common Level 2 bullets
            '■', '□', '▪️', '▣', '▤', '⬧', '➔', '⬛', 'o', '✓', '✔', '❖'  # Common Level 3 bullets
        ]

        current_level = 0  # Default starting level for top-level bullets
        previous_text = ""  # Variable to track the previous bullet point text
        previous_indentation = 0  # To track the indentation level of the previous bullet point

        try:
            for page_num, page in enumerate(self.doc, 1):
                blocks = page.get_text("dict")["blocks"]

                for block in blocks:
                    # print(f'Block: {block}')
                    x0 = block['bbox'][0]
                    # print(x0)
                    indentation_level = int(x0)
                    # print(x0)
                    if "lines" in block:
                        for i, line in enumerate(block["lines"]):
                            for span in line["spans"]:
                                # text = span["text"]
                                text = span["text"].strip()
                                # print(text)
                                # Check if line starts with a recognized bullet point
                                if text and text[0] in set(recognized_bullet_symbols):
                                    logger.info(f"Found bullet point on page {page_num}: {text}")
                                    
                                    # Skip if the bullet point is not properly formatted (needs space after symbol)
                                    if len(text) > 1 and text[1] != ' ':
                                        continue
                                    
                                    # print(text)
                                    # Count leading spaces (indentation) to determine bullet depth
                                    # indentation_level = len(text) - len(text.lstrip())

                                    # Determine if the current bullet is a child of the previous one
                                    if indentation_level > previous_indentation:
                                        # Indentation level increased, meaning it's a child of the previous bullet point
                                        current_level = min(current_level + 1, 3)  # Max level is 3
                                        logger.info(f"Bullet point is a child of the previous bullet: {text}")
                                    elif indentation_level < previous_indentation:
                                        # Indentation level decreased, meaning it's at the same level as a previous bullet
                                        current_level = 1  # Reset to level 1 for a new root bullet
                                        logger.info(f"Bullet point is not at the same level as a previous bullet: {text}")
                                    else:
                                        # Same indentation level as the previous bullet, so it stays at the same level
                                        current_level = current_level
                                        logger.info(f"Bullet point is at the same level as the previous bullet: {text}")
                                        
                                    # Validate the bullet point's format
                                    if text[0] not in bullet_patterns[current_level]:
                                        error_msg = {
                                            "page": page_num,
                                            "text": text,
                                            "incorrect_symbol": text[0],
                                            "level": current_level,
                                            "expected_symbol": bullet_patterns[current_level]
                                        }
                                        logger.error(f"Formatting error: {error_msg}")
                                        errors.append(error_msg)
                                    else:
                                        logger.info(f"Correct bullet format found for level {current_level}: {text[0]}")
                                        
                                    # Update previous bullet point for next iteration
                                    print(text, indentation_level)
                                    previous_text = text
                                    previous_indentation = indentation_level
                            
        except Exception as e:
            error_msg = f"Error in bullet points check: {str(e)}"
            logger.error(error_msg)
            raise

        logger.info(f"Completed bullet points check. Found {len(errors)} errors")
        return errors

    def review_document(self) -> Dict[str, List[str]]:
        """Perform complete document format review"""
        logger.info("Starting complete PDF document review")
        try:
            results = {
                'heading_numbering': self.check_heading_numbering(),
                'page_margins': self.check_page_margins(),
                'bullet_points': self.check_bullet_points()
            }
            
            total_errors = sum(len(issues) for issues in results.values())
            logger.info(f"PDF document review completed. Total errors found: {total_errors}")
            return results
            
        except Exception as e:
            error_msg = f"Failed to complete PDF document review: {str(e)}"
            logger.error(error_msg)
            raise
        finally:
            self.doc.close()