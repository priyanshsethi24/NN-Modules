from abc import ABC, abstractmethod
from docx import Document
import PyPDF2
import re
from typing import Set, Dict, List, Optional
from urllib.parse import urlparse
import os
from pdf2docx import Converter

class LinkExtractor(ABC):
    """Abstract base class for document link extractors"""
    @abstractmethod
    def extract_links(self, file_path: str) -> Dict[str, Dict]:
        """Extract external links from a document with their page numbers and display text"""
        pass

    def is_external_link(self, url: str) -> bool:
        """Check if a URL is an external link"""
        try:
            parsed = urlparse(url)
            return bool(parsed.scheme and parsed.netloc) or parsed.scheme in {'file', 'smb', 'nfs'}
        except Exception:
            return False

class DocxLinkExtractor(LinkExtractor):
    """DOCX file link extractor"""
    def extract_links(self, file_path: str) -> Dict[str, Dict]:
        doc = Document(file_path)
        links = {}
        
        # Calculate pages based on paragraph count
        total_paragraphs = len(doc.paragraphs)
        paragraphs_per_page = 40  # Approximate number of paragraphs per page
        
        # Extract hyperlinks from paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            current_page = (i // paragraphs_per_page) + 1
            
            # First try to get hyperlinks with their text
            for hyperlink in paragraph._element.findall(".//w:hyperlink", {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                rId = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if rId:
                    # Get the relationship target (URL)
                    relationship = doc.part.rels[rId]
                    if relationship.reltype.endswith('/hyperlink'):
                        url = relationship._target
                        
                        # Get the display text
                        texts = []
                        for run in hyperlink.findall(".//w:t", {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            if run.text:
                                texts.append(run.text)
                        display_text = ''.join(texts)
                        
                        if not display_text:
                            display_text = "Link"
                            
                        if self.is_external_link(url):
                            if url not in links:
                                links[url] = {
                                    "display_text": display_text,
                                    "pages": []
                                }
                            if current_page not in links[url]["pages"]:
                                links[url]["pages"].append(current_page)
        
        return links

class PdfLinkExtractor(LinkExtractor):
    """PDF file link extractor"""
    def extract_links(self, file_path: str) -> Dict[str, Dict]:
        links = {}
        temp_files = []
        
        try:
            # First get the original PDF for page number reference
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pdf_text_by_page = {}
                
                # Store text content of each PDF page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    pdf_text_by_page[page_num] = page.extract_text()

            # Convert PDF to DOCX for better text extraction
            docx_path = file_path.rsplit('.', 1)[0] + '_temp.docx'
            temp_files.append(docx_path)
            
            try:
                # Convert PDF to DOCX
                cv = Converter(file_path)
                cv.convert(docx_path)
                cv.close()
                
                # Load the DOCX document
                doc = Document(docx_path)
                
                # Process each paragraph
                for paragraph in doc.paragraphs:
                    # Get hyperlinks from paragraph's XML element
                    root = paragraph._element
                    for hyperlink in root.findall('.//w:hyperlink', 
                        {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        # Get relationship id
                        rel_id = hyperlink.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'
                        )
                        
                        if rel_id and rel_id in paragraph.part.rels:
                            # Get URL from relationship
                            url = paragraph.part.rels[rel_id]._target
                            
                            if self.is_external_link(url):
                                # Get text from the hyperlink
                                link_text = ''.join(
                                    t.text for t in hyperlink.findall(
                                        './/w:t', 
                                        {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                                    )
                                )
                                
                                # Use actual link text or fallback
                                display_text = link_text if link_text else "Link"
                                
                                # Find the actual page number from PDF
                                page_number = self._find_page_number(pdf_text_by_page, display_text)
                                
                                if url not in links:
                                    links[url] = {
                                        "display_text": display_text,
                                        "pages": []
                                    }
                                if page_number:  # Only add page number if found
                                    links[url]["pages"].append(page_number)
                                else:
                                    links[url]["pages"].append("Not Found")  # Add "Not Found" instead of a page number
                
            except Exception as e:
                print(f"Error in DOCX conversion: {str(e)}")
                # Fallback to original PDF extraction if conversion fails
                return self._extract_from_pdf(file_path)
                
        finally:
            # Clean up temporary files
            for temp_file in temp_files:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
        
        return links

    def _find_page_number(self, pdf_text_by_page: Dict[int, str], search_text: str) -> Optional[int]:
        """Find the page number where the text appears in the PDF"""
        for page_num, text in pdf_text_by_page.items():
            if search_text in text:
                return page_num
        return None  # Return None instead of defaulting to page 1

    def _extract_from_pdf(self, file_path: str) -> Dict[str, Dict]:
        """Fallback method to extract links directly from PDF"""
        links = {}
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages, 1):
                    if '/Annots' in page:
                        annotations = page['/Annots']
                        for annotation in annotations:
                            if isinstance(annotation, PyPDF2.generic.IndirectObject):
                                annotation = annotation.get_object()
                            
                            try:
                                if annotation.get('/Subtype') == '/Link' and '/A' in annotation:
                                    action = annotation['/A']
                                    if isinstance(action, PyPDF2.generic.IndirectObject):
                                        action = action.get_object()
                                    
                                    if '/URI' in action:
                                        uri = action['/URI']
                                        if isinstance(uri, PyPDF2.generic.IndirectObject):
                                            uri = uri.get_object()
                                        
                                        if self.is_external_link(uri):
                                            # Default to "Link" if we can't get the text
                                            if uri not in links:
                                                links[uri] = {
                                                    "display_text": "Link",
                                                    "pages": []
                                                }
                                            if page_num not in links[uri]["pages"]:
                                                links[uri]["pages"].append(page_num)
                            except Exception as e:
                                print(f"Error processing annotation: {str(e)}")
                                continue
            
            return links
            
        except Exception as e:
            print(f"Error processing PDF: {str(e)}")
            raise

class LinkExtractorFactory:
    """Factory for creating appropriate link extractors"""
    _extractors = {
        '.docx': DocxLinkExtractor,
        '.pdf': PdfLinkExtractor
    }
    
    @classmethod
    def get_extractor(cls, file_extension: str) -> LinkExtractor:
        extractor_class = cls._extractors.get(file_extension.lower())
        if not extractor_class:
            raise ValueError(f"Unsupported file format: {file_extension}")
        return extractor_class()

class LinkExtractorService:
    """Service class for extracting links from documents"""
    def __init__(self, factory: LinkExtractorFactory):
        self.factory = factory

    def extract_links(self, file_path: str) -> Dict[str, Dict]:
        """Extract links from a document file"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1]
        extractor = self.factory.get_extractor(file_extension)
        return extractor.extract_links(file_path)

# ... rest of the code from Module6.py ... 